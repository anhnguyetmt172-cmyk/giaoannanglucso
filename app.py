import streamlit as st
import google.generativeai as genai
from PIL import Image
import tempfile
import os
import io
import re

# ===== WORD (python-docx) =====
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== 1. CẤU HÌNH TRANG =====
st.set_page_config(
    page_title="Trợ lý Soạn Giáo án NLS",
    page_icon="📘",
    layout="centered"
)

FILE_KHUNG_NANG_LUC = "khungnanglucso.pdf"

# ===== 2. HÀM XỬ LÝ WORD =====
def add_formatted_text(paragraph, text):
    """In đậm **text**, font Times New Roman 14"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)


def create_doc_stable(content, ten_bai, lop):
    doc = Document()

    # --- A4 + LỀ CHUẨN ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.3

    # --- TIÊU ĐỀ ---
    title = doc.add_heading(f"KẾ HOẠCH BÀI DẠY: {ten_bai.upper()}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)

    p_lop = doc.add_paragraph(f"Lớp: {lop}")
    p_lop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lop.runs[0].bold = True

    doc.add_paragraph("-" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- XỬ LÝ NỘI DUNG ---
    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if line.startswith("#"):
            line = line.replace("#", "").strip()

        # ===== BẢNG MARKDOWN =====
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            rows = [r for r in table_lines if "---" not in r]
            if rows:
                cols = len(rows[0].split("|")) - 2
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"

                for r_idx, row in enumerate(rows):
                    cells = row.split("|")[1:-1]
                    for c_idx, cell_text in enumerate(cells):
                        cell = table.cell(r_idx, c_idx)
                        cell._element.clear_content()

                        sub_lines = cell_text.replace("<br>", "\n").split("\n")
                        for sub in sub_lines:
                            if not sub.strip():
                                continue
                            p = cell.add_paragraph()
                            if r_idx == 0:
                                run = p.add_run(sub.replace("**", ""))
                                run.bold = True
                                run.font.name = "Times New Roman"
                                run.font.size = Pt(14)
                            else:
                                add_formatted_text(p, sub.strip())
            continue

        if not line:
            i += 1
            continue

        # ===== MỤC I, II, III =====
        if re.match(r'^(I\.|II\.|III\.|IV\.|V\.)', line):
            p = doc.add_paragraph(line.replace("**", ""))
            p.runs[0].bold = True
            p.runs[0].font.name = "Times New Roman"
            p.runs[0].font.size = Pt(14)

        # ===== GẠCH ĐẦU DÒNG THỦ CÔNG =====
        elif line.startswith("- "):
            p = doc.add_paragraph()
            add_formatted_text(p, f"• {line[2:].strip()}")

        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)

        i += 1

    return doc


# ===== 3. GIAO DIỆN =====
st.markdown("""
<style>
[data-testid="stAppViewContainer"] { background-color: #f4f6f9; }
.lesson-plan-paper {
    background: white;
    padding: 40px;
    border: 1px solid #ccc;
    font-family: 'Times New Roman';
    font-size: 14pt;
    line-height: 1.5;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div style="text-align:center">
<h2>📘 TRỢ LÝ SOẠN GIÁO ÁN TỰ ĐỘNG</h2>
<p><i>Nguyễn Thị Ánh Nguyệt – Trường Tiểu học Ngọc Đường</i></p>
</div>
""", unsafe_allow_html=True)

# ===== 4. API KEY =====
if "GEMINI_API_KEY" in st.secrets:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
else:
    api_key = st.text_input("Nhập Gemini API Key:", type="password")
    if api_key:
        genai.configure(api_key=api_key)

# ===== 5. INPUT =====
lop = st.text_input("Lớp:", "Lớp 4")
ten_bai = st.text_input("Tên bài học:")
noidung = st.text_area("Ghi chú / yêu cầu thêm:")

uploaded_files = st.file_uploader(
    "Tải ảnh hoặc PDF bài học:",
    type=["jpg", "png", "pdf"],
    accept_multiple_files=True
)

# ===== 6. XỬ LÝ =====
if st.button("🚀 SOẠN GIÁO ÁN"):
    with st.spinner("AI đang soạn giáo án..."):
        model = genai.GenerativeModel("gemini-2.5-flash-lite-preview-09-2025")
        prompt = noidung if noidung else "Soạn giáo án theo CT GDPT 2018, Công văn 2345."
        response = model.generate_content(prompt)

        st.markdown("### 📄 KẾT QUẢ")
        st.markdown(f"<div class='lesson-plan-paper'>{response.text}</div>", unsafe_allow_html=True)

        doc = create_doc_stable(response.text, ten_bai, lop)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        st.download_button(
            "⬇️ TẢI FILE WORD A4",
            data=buf,
            file_name=f"GiaoAn_{ten_bai}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ===== CHÂN TRANG =====
st.markdown("---")
st.markdown(
    "<div style='text-align:center;color:#666'>© 2025 – Nguyễn Thị Ánh Nguyệt – TH Ngọc Đường</div>",
    unsafe_allow_html=True
)
